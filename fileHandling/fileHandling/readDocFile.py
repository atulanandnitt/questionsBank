

import docx2txt

#method #1
def readDoc():

    # extract text
    text = docx2txt.process("basic_practice_notes.docx")
    print(text)
    # extract text and write images in /tmp/img_dir
    text2 = docx2txt.process("basic_practice_notes.docx", "/tmp/img_dir")
    print("data for text2" ,text2)
#readDoc()


#method #2
import docx#pip install python-docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        #print("para",para.text)
        fullText.append(para.text)
    return '\n'.join(fullText)

filename="basic_practice_notes.docx"
print(getText(filename))


from docx import Document
def writeDocFile():
    document = Document()
    document.add_paragraph('Intense quote')

    document.save('basic_output.docx')

writeDocFile()




print("advanced docx file read write")
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

#document.add_picture('monty-truth.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

document.add_page_break()

document.save('demo.docx')